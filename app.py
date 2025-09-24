from gensim.models.doc2vec import Doc2Vec
model = Doc2Vec.load('models/d2v.model')
from flask import Flask, render_template, request, redirect, session, flash, url_for, send_file, abort
import os
import re
import psycopg2
import psycopg2.extras
from psycopg2 import sql
import pdfplumber
import spacy
import io
import base64
import time
import json
import html
import matplotlib
import matplotlib.pyplot as plt
import fitz  # PyMuPDF
import datetime
import random
from docx import Document
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
from Sastrawi.StopWordRemover.StopWordRemoverFactory import StopWordRemoverFactory
from similarity_utils import get_stemmed_tokens_from_similar_sentences

from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from sklearn.feature_extraction.text import TfidfVectorizer
# Reportlab
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
from reportlab.lib import colors
from reportlab.lib.colors import yellow, green, blue, red, orange
from preprocess import preprocess_sentences
from model_loader import model  # jika model dimuat dari file luar
from preprocess import preprocess_text
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from pdf_utils import create_combined_exum_highlighted_pdf
# App-specific
from auth import role_required
from app.helpers.file_helpers import extract_text_from_file
# from train_model import prepare_and_train_model
from update_vectors import update_document_vectors
from similarity_utils import (
    split_and_clean_sentences,
    embed_sentences,
    get_top_similar_pairs,
    highlight_matches
)


# Set backend non-interaktif untuk matplotlib
matplotlib.use('Agg')


app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.secret_key = 'secret_key'

# Load model with error handling
try:
    model = Doc2Vec.load('models/d2v.model')
except Exception as e:
    print(f"Error loading model: {e}")

def generate_highlight_filename(filename):
    return filename.replace('.', '_') + "_highlight.json"

def generate_unique_filename(prefix="skripsi_highlight", ext=".pdf"):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    rand = random.randint(1000, 9999)
    return f"{prefix}_{timestamp}_{rand}{ext}"

def get_db_connection():
    """Handles database connection pooling and querying."""
    conn = psycopg2.connect(
        host="localhost",
        database="db_similarity",
        user="postgres",
        password="admin",
        port="5432"
        
    )
    return conn

def highlight_pdf_based_on_tokens(pdf_path, output_path, similar_tokens):
    print("Input PDF Path:", pdf_path)
    print("Output PDF Path:", output_path)

    factory = StemmerFactory()
    stemmer = factory.create_stemmer()

    # Preprocess similar_tokens (lower, stem, stopword removal)
    stemmed_tokens = set(
        stemmer.stem(t.lower()) 
        for t in similar_tokens 
        if t.lower() not in INDONESIAN_STOPWORDS and t.isalpha()
    )
    print("Stemmed Similar Tokens:", stemmed_tokens)

    matched_words = set()

    doc = fitz.open(pdf_path)

    for page_num, page in enumerate(doc):
        words = page.get_text("words")
        for w in words:
            word = w[4].lower()
            if word in INDONESIAN_STOPWORDS or not word.isalpha():
                continue  # skip stopwords & non-alphabet
            stemmed = stemmer.stem(word)
            if stemmed in stemmed_tokens and len(word) > 2:
                rect = fitz.Rect(w[0], w[1], w[2], w[3])
                page.add_highlight_annot(rect)
                matched_words.add(word)

    doc.save(output_path)
    doc.close()
    print(f"✅ Saved highlighted PDF to: {output_path}")
    print(f"✨ Total words matched: {len(matched_words)}")
    print(f"📝 Sample matched words: {list(matched_words)[:10]}")

def extract_keywords(text, top_n=30):
    from Sastrawi.Stemmer.StemmerFactory import StemmerFactory

    factory = StemmerFactory()
    stemmer = factory.create_stemmer()

    # 1. Pre-cleaning: lowercase & remove symbols
    text_cleaned = re.sub(r'[^a-z\s]', '', text.lower())  # hanya huruf kecil & spasi

    # 2. TF-IDF vectorizer
    vectorizer = TfidfVectorizer(stop_words=list(INDONESIAN_STOPWORDS), max_features=1000)
    tfidf_matrix = vectorizer.fit_transform([text_cleaned])

    # 3. Urutkan berdasarkan skor
    scores = zip(vectorizer.get_feature_names_out(), tfidf_matrix.toarray()[0])
    sorted_words = sorted(scores, key=lambda x: x[1], reverse=True)

    # 4. Ambil top-N keyword dan stem-kan
    keywords = [word for word, _ in sorted_words[:top_n]]
    stemmed_keywords = [stemmer.stem(word) for word in keywords]

    return stemmed_keywords


def preprocess_text(text):
    from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
    import re

    # Inisialisasi stemmer
    factory = StemmerFactory()
    stemmer = factory.create_stemmer()

    # Join kalau input berupa list
    if isinstance(text, list):
        text = " ".join([str(item[0]) if isinstance(item, tuple) else str(item) for item in text])

    # Lowercase & bersihkan karakter non-huruf
    text = text.lower()
    text = re.sub(r'[^a-z\s]', '', text)

    # Tokenisasi
    tokens = text.split()

    # Hapus stopwords & token pendek
    filtered = [word for word in tokens if word not in INDONESIAN_STOPWORDS and len(word) > 2]

    # Stemming
    stemmed_tokens = [stemmer.stem(word) for word in filtered]

    return stemmed_tokens


# Menarik dokumen dari database
def fetch_documents_from_db():
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Query untuk mengambil semua dokumen dari database
    cur.execute("SELECT id, title, file_text FROM documents")
    documents = cur.fetchall()

    cur.close()
    conn.close()

    return documents
def calculate_similarity(user_doc, db_docs):
    factory = StemmerFactory()
    stemmer = factory.create_stemmer()

    # Token & keywords user
    user_tokens = preprocess_text(user_doc)
    if not user_tokens:
        return []  # kalau kosong, langsung keluar
    user_vector = model.infer_vector(user_tokens)

    # TF-IDF keywords + stemming
    user_keywords = extract_keywords(user_doc, top_n=50)
    user_stemmed = set(stemmer.stem(token.lower()) for token in user_keywords)

    # Vektor kalimat user
    user_sentences = split_and_clean_sentences(user_doc)
    user_vecs = embed_sentences(user_sentences, model)

    similarities = []

    for db_doc in db_docs:
        db_text = db_doc['file_text']
        db_tokens = preprocess_text(db_text)
        if not db_tokens:
            continue
        db_vector = model.infer_vector(db_tokens)

        # Cosine similarity dokumen
        try:
            cosine_score = float(cosine_similarity([user_vector], [db_vector])[0][0])
        except ValueError:
            cosine_score = 0.0

        # Stem token DB untuk deteksi token mirip
        db_stemmed = set(stemmer.stem(token.lower()) for token in db_tokens)
        similar_tokens = db_stemmed.intersection(user_stemmed)

        mark_count = len(similar_tokens)
        db_token_count = len(db_stemmed)
        token_match_percent = (mark_count / db_token_count) * 100 if db_token_count else 0

        # Kalimat & vektor untuk top 3 kalimat paling mirip
        db_sentences = split_and_clean_sentences(db_text)
        db_vecs = embed_sentences(db_sentences, model)
        top_pairs = get_top_similar_pairs(user_sentences, user_vecs, db_sentences, db_vecs)

        # Skor kombinasi
        combined_score = (0.7 * cosine_score) + (0.3 * (token_match_percent / 100))

        similarities.append({
            'id': db_doc['id'],
            'title': db_doc['title'],
            'cosine_score': cosine_score,
            'token_match_percent': token_match_percent,
            'combined_score': combined_score,
            'similar_tokens': similar_tokens,
            'top_sentence_pairs': top_pairs
        })

    return sorted(similarities, key=lambda x: x['combined_score'], reverse=True)[:5]


def remove_html_tags(text):
    """Menghapus semua tag HTML seperti <span>, <font>, dll."""
    return re.sub(r'<[^>]+>', '', text)
# Fungsi utama untuk mendapatkan 5 dokumen dengan skor similarity tertinggi
def get_top_similar_documents(exum_text):
    db_documents = fetch_documents_from_db()
    top_5_docs = calculate_similarity(exum_text, db_documents)
    return top_5_docs

def init_db():
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    cur.execute("""
        SELECT h.id, d.title, h.similarity_score, h.uploaded_file_text
        FROM history h
        JOIN documents d ON h.document_id = d.id
        WHERE h.result_file_path IS NOT NULL
        ORDER BY h.id DESC
    """)

    conn.commit()
    cur.close()
    conn.close()

def save_result_pdf(user_id, file_name, original_text, similar_tokens, history_id=None):
    result_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'results', str(user_id))
    os.makedirs(result_dir, exist_ok=True)

    base_name = os.path.splitext(file_name)[0]
    exum_pdf_path = os.path.join(result_dir, f"{base_name}_exum_combined.pdf")

    # Hanya simpan Exum highlight
    create_highlighted_pdf(original_text, similar_tokens, exum_pdf_path, user_id=user_id, filename_hint=file_name)


    return exum_pdf_path


def get_file_path_by_history_id(history_id):
    docs = get_documents_from_db()
    for doc in docs:
        if doc['id'] == history_id:
            return doc.get('file_path')  # atau 'filename' tergantung key yang kamu simpan
    return None


@app.route('/')
def landing_page():
    return render_template('landing_page.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM users WHERE username = %s", (username,))
        user = cur.fetchone()
        cur.close()
        conn.close()

        if user:
            stored_password = user['password']

            # 1️⃣ Jika password tersimpan sebagai hash (pbkdf2/scrypt/argon2 dll)
            if stored_password.startswith("pbkdf2:") or stored_password.startswith("scrypt:"):
                if check_password_hash(stored_password, password):
                    session['username'] = user['username']
                    session['role'] = user['role']
                    session['user_id'] = user['id']
                    return redirect(url_for('home' if user['role'] == 'admin' else 'dashboard'))

            # 2️⃣ Jika password tersimpan plain text
            elif stored_password == password:
                session['username'] = user['username']
                session['role'] = user['role']
                session['user_id'] = user['id']
                return redirect(url_for('home' if user['role'] == 'admin' else 'dashboard'))

        flash("Username atau password salah", "danger")
    return render_template("landing_page.html")

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        role = request.form['role']
        
        # Buat hash password
        hashed_password = generate_password_hash(password)

        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        try:
            cur.execute(
                "INSERT INTO users (username, password, role) VALUES (%s, %s, %s)",
                (username, hashed_password, role)
            )
            conn.commit()
            # Registrasi berhasil, alihkan ke login
            flash('Registrasi berhasil! Silakan login.', 'success')
            return redirect(url_for('login'))
        except psycopg2.IntegrityError:
            conn.rollback()
            # Pengguna sudah ada
            flash('Username sudah ada. Silakan gunakan nama lain.', 'danger')
            return render_template('landing_page')
        finally:
            cur.close()
            conn.close()

    return render_template('landing_page')

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/login')
@app.route('/dashboard')
@role_required(['user'])
def dashboard():
    user_id = session.get('user_id')

    if not user_id:
        return redirect('/login')

    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # 1. Ambil jumlah total dokumen pengguna
    cur.execute("SELECT COUNT(id) AS total FROM history WHERE user_id = %s", (user_id,))
    total_documents = cur.fetchone()['total']

    # 2. Ambil skor rata-rata
    cur.execute("SELECT AVG(similarity_score) AS avg_score FROM history WHERE user_id = %s", (user_id,))
    avg_score_raw = cur.fetchone()['avg_score']
    avg_score = round(avg_score_raw, 2) if avg_score_raw is not None else 0.0

    # 3. Ambil dokumen terakhir yang diunggah
    cur.execute("""
        SELECT uploaded_file_name, similarity_score
        FROM history
        WHERE user_id = %s
        ORDER BY created_at DESC
        LIMIT 1
    """, (user_id,))
    last_doc = cur.fetchone()
    last_doc_name = last_doc['uploaded_file_name'] if last_doc else 'N/A'
    last_doc_score = round(last_doc['similarity_score'], 2) if last_doc and last_doc['similarity_score'] is not None else 'N/A'

    # 4. Ambil 5 riwayat terbaru untuk tabel
    cur.execute("""
        SELECT id, uploaded_file_name, similarity_score, created_at
        FROM history
        WHERE user_id = %s
        ORDER BY created_at DESC
        LIMIT 5
    """, (user_id,))
    recent_history = cur.fetchall()

    cur.close()
    conn.close()

    return render_template(
        'dashboard.html',
        username=session['username'],
        role=session['role'],
        active_page='dashboard',
        total_documents=total_documents,
        avg_score=avg_score,
        last_document_name=last_doc_name,
        last_document_score=last_doc_score,
        recent_history=recent_history
    )



@app.route('/home')
@role_required(['admin'])
def home():
    if 'username' not in session:
        return redirect('/login')

    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Statistik jumlah dokumen per user
    cur.execute("""
        SELECT u.username, COUNT(h.id) AS count
        FROM users u 
        LEFT JOIN history h ON u.id = h.user_id 
        GROUP BY u.username
    """)
    doc_stats = cur.fetchall()

    # Total users
    cur.execute("SELECT COUNT(*) FROM users")
    total_users = cur.fetchone()['count']

    # Total documents
    cur.execute("SELECT COUNT(*) FROM documents")
    total_documents = cur.fetchone()['count']

    # Upload hari ini
    cur.execute("SELECT COUNT(*) FROM history WHERE DATE(created_at) = CURRENT_DATE")
    today_uploads = cur.fetchone()['count']

    # Aktivitas terbaru (5 terakhir)
    cur.execute("""
        SELECT u.username, h.uploaded_file_name AS file_name, h.created_at
        FROM history h
        JOIN users u ON u.id = h.user_id
        ORDER BY h.created_at DESC
        LIMIT 5
    """)
    recent_activities = cur.fetchall()

    cur.close()
    conn.close()

    # Data untuk grafik
    usernames = [row['username'] for row in doc_stats]
    doc_counts = [row['count'] for row in doc_stats]

    return render_template(
        "home.html",
        username=session['username'],
        role=session['role'],
        total_users=total_users,
        total_documents=total_documents,
        today_uploads=today_uploads,
        recent_activities=recent_activities,
        usernames=usernames,
        doc_counts=doc_counts,
        active_page='dashboard'
    )


@app.route('/admin/upload', methods=['GET', 'POST'])
def admin_upload():
    if 'role' not in session or session['role'] != 'admin':
        return redirect('/login')

    if request.method == 'POST':
        title = request.form['title']
        file = request.files['file']
        document_type = request.form.get('document_type', 'skripsi')  # default ke skripsi

        if file and allowed_file(file.filename):
            filename = f"{int(time.time())}_{secure_filename(file.filename)}"
            admin_folder = os.path.join(app.config['UPLOAD_FOLDER'], 'admin_documents')
            os.makedirs(admin_folder, exist_ok=True)
            filepath = os.path.join(admin_folder, filename)
            file.save(filepath)

            start_time = time.time()
            file_text = extract_text_from_file(filepath)
            print("Text extraction took:", time.time() - start_time, "seconds")

            file_text = file_text.encode("utf-8", "replace").decode("utf-8")
            file_text = file_text.replace("\x00", "")

            tokens = preprocess_text(file_text)
            vector = model.infer_vector(tokens)
            vector_bytes = np.array(vector).tobytes()

            conn = get_db_connection()
            cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
            try:
                print("Uploading:", title, filename)
                cur.execute(
                    """
                    INSERT INTO documents (title, file_name, file_text, vector, document_type)
                    VALUES (%s, %s, %s, %s, %s)
                    """,
                    (title, filename, file_text, vector_bytes, document_type)
                )
                conn.commit()
                print("Insert success")
            except Exception as e:
                conn.rollback()
                print("Insert failed:", e)
            finally:
                cur.close()
                conn.close()

            flash('Dokumen berhasil di-upload!', 'success')
            return redirect('/admin/upload')

    return render_template(
        'admin/upload.html',
        active_page='upload'  
    )
    

def allowed_file(filename):
    allowed_extensions = ['pdf', 'docx']
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

@app.route('/admin/documents')
@role_required(['admin'])
def admin_documents():
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT id, title, file_name, document_type FROM documents ORDER BY id DESC")
    documents = cur.fetchall()
    print("Dokumen yang ditemukan:", documents)  # DEBUG
    cur.close()
    conn.close()

    return render_template(
        'admin/documents.html',
        documents=documents,
        active_page='documents'   # ✅ Tambahkan ini
    )


@app.route('/view/<int:history_id>')
def view_result(history_id):
    # Bisa diarahkan ke admin_view_result atau user_view_result
    user_id = session.get('user_id')
    return redirect(url_for('admin_view_result', user_id=user_id, history_id=history_id))

@app.route('/user/documents')
@role_required(['user'])
def user_documents():
    user_id = session.get('user_id')
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Query to get documents with similarity_score
    cur.execute("""
        SELECT h.id, h.similarity_score, d.title, h.created_at, h.uploaded_file_name, h.uploaded_file_text
        FROM history h
        JOIN documents d ON h.document_id = d.id
        WHERE h.user_id = %s AND h.result_file_path IS NOT NULL
        ORDER BY h.id DESC
    """, (user_id,))

    documents = cur.fetchall()

    # Process similarity score and date formatting
    for doc in documents:
        if doc['similarity_score'] is not None:
            doc['average_score'] = round(min(doc['similarity_score'], 100), 2)

        else:
            doc['average_score'] = 0.0
    

        doc['created_at'] = doc['created_at'].strftime('%Y-%m-%d')


    cur.close()
    conn.close()

    return redirect(url_for('history_user_view', user_id=session['user_id']))

# Halaman Kelola User
@app.route('/manage_users')
@role_required(['admin'])
def manage_users():
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT id, username, role FROM users ORDER BY id ASC")
    users = cur.fetchall()
    cur.close()
    conn.close()
    return render_template("manage_users.html", users=users, active_page='manage_users')

@app.route('/add_user', methods=['GET', 'POST'])
@role_required(['admin'])
def add_user():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        role = request.form['role']

        hashed_password = generate_password_hash(password)

        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        try:
            cur.execute(
                "INSERT INTO users (username, password, role) VALUES (%s, %s, %s)",
                (username, hashed_password, role)
            )
            conn.commit()
            flash("User berhasil ditambahkan!", "success")  # alert hijau
            return redirect(url_for('manage_users'))  # redirect ke daftar user
        except psycopg2.errors.UniqueViolation:
            conn.rollback()
            flash('Username sudah ada. Silakan gunakan nama lain.', 'danger')  # alert merah
            return redirect(url_for('add_user'))  # kembali ke form
        finally:
            cur.close()
            conn.close()

    return render_template("add_user.html", active_page='manage_users')

# Edit User
@app.route('/edit_user/<int:user_id>', methods=['GET', 'POST'])
@role_required(['admin'])
def edit_user(user_id):
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    if request.method == 'POST':
        username = request.form['username']
        role = request.form['role']
        cur.execute("UPDATE users SET username=%s, role=%s WHERE id=%s", 
                    (username, role, user_id))
        conn.commit()
        cur.close()
        conn.close()
        flash("User berhasil diperbarui!", "success")
        return redirect(url_for('manage_users'))

    cur.execute("SELECT * FROM users WHERE id=%s", (user_id,))
    user = cur.fetchone()
    cur.close()
    conn.close()
    return render_template("edit_user.html", user=user, active_page='manage_users')

# Hapus User
@app.route('/delete_user/<int:user_id>', methods=['POST'])
@role_required(['admin'])
def delete_user(user_id):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("DELETE FROM users WHERE id = %s", (user_id,))
    conn.commit()
    cur.close()
    conn.close()
    flash("User berhasil dihapus!", "success")
    return redirect(url_for('manage_users'))



@app.route('/hasil')
def hasil():
    if 'uploaded_file' not in session:
        flash('Tidak ada dokumen yang di-upload.', 'danger')
        return redirect('/upload')

    user_id = session.get("user_id", "guest")
    filename = session['uploaded_file']
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], 'user_documents', filename)

    with open(filepath, 'rb') as f:
        exum_text = extract_text_from_file(f)

    # Ambil dokumen dari DB
    db_documents = fetch_documents_from_db()

    # 🔥 Highlight + Similarity
    highlighted_path = generate_highlighted_exum_pdf(user_id, exum_text, db_documents)

    return send_file(highlighted_path, as_attachment=True)



@app.route('/user/delete/<int:history_id>')
@role_required(['user'])
def delete_history_entry(history_id):
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Ambil file path untuk dihapus dari sistem
    # Mengakses result menggunakan indeks numerik
    cur.execute("SELECT result_file_path FROM history WHERE id = %s", (history_id,))
    result = cur.fetchone()

    # Pastikan menggunakan result[0], karena hanya ada satu kolom yang diambil
    if result and result[0] and os.path.exists(result[0]):
        return send_file(result[0], mimetype='application/pdf', as_attachment=False)


    # Hapus entri dari database
    cur.execute("DELETE FROM history WHERE id = %s", (history_id,))
    conn.commit()
    cur.close()
    conn.close()

    flash('Histori berhasil dihapus.', 'success')
    return redirect('/user/documents')

@app.route('/user/download/<int:history_id>')
def download_user_result(history_id):
    conn = psycopg2.connect(
        "dbname=db_similarity user=postgres password=admin host=localhost"
    )
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT result_file_path FROM history WHERE id = %s", (history_id,))
    result = cur.fetchone()
    cur.close()
    conn.close()

    if result and result['result_file_path'] and os.path.exists(result['result_file_path']):
        return send_file(result[0], mimetype='application/pdf', as_attachment=True, download_name=f"Hasil_Similaritas_{history_id}.pdf")
    abort(404, description="File tidak ditemukan.")

@app.route('/history-user')
@role_required(['user'])
def history_user_view():
    user_id = session.get('user_id')
    if not user_id:
        flash("Anda belum login.", "error")
        return redirect('/login')

    # Ambil parameter filter dari query string
    title = request.args.get('title')
    min_score = request.args.get('min_score')
    max_score = request.args.get('max_score')

    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Query history yang JOIN ke tabel documents
    query = """
        SELECT 
            MIN(h.id) AS id,
            h.user_id,
            h.uploaded_file_name,
            h.uploaded_file_text,
            AVG(h.similarity_score) AS average_score,
            h.created_at
        FROM history h
        WHERE h.user_id = %s
        GROUP BY h.uploaded_file_name, h.uploaded_file_text, h.created_at, h.user_id
    """

    params = [user_id]

    if title:
        query += " AND h.uploaded_file_name ILIKE %s"
        params.append(f"%{title}%")
    if min_score:
        query += " AND h.similarity_score >= %s"
        params.append(min_score)
    if max_score:
        query += " AND h.similarity_score <= %s"
        params.append(max_score)

    query += " ORDER BY h.created_at DESC"
    cur.execute(query, tuple(params))
    documents = cur.fetchall()

    cur.close()
    conn.close()

    return render_template(
        'user/documents.html',
        documents=documents,
        active_page='history'   # ✅ biar sidebar aktif
    )


@app.route('/history-user/<int:history_id>')
@role_required(['user'])
def history_user_view_detail(history_id):
    user_id = session.get('user_id')
    if not user_id:
        flash("Anda belum login.", "error")
        return redirect('/login')

    title = request.args.get('title')
    min_score = request.args.get('min_score')
    max_score = request.args.get('max_score')

    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Ambil waktu dan nama file upload-nya
    cur.execute(
        "SELECT created_at, uploaded_file_name FROM history WHERE user_id = %s AND id = %s",
        (user_id, history_id)
    )
    result = cur.fetchone()
    if not result:
        flash("Data tidak ditemukan.", "error")
        return redirect('/dashboard')

    created_at = result['created_at']
    uploaded_file_name = result['uploaded_file_name']

    # Query utama
    query = """
        SELECT 
            h.id AS history_id,
            h.similarity_score,
            h.uploaded_file_name,
            h.uploaded_file_text,
            d.file_name AS compared_file_name
        FROM history h
        JOIN documents d ON h.document_id = d.id
        WHERE h.created_at = %s AND h.uploaded_file_name = %s
    """
    params = [created_at, uploaded_file_name]

    if title:
        query += " AND h.uploaded_file_name ILIKE %s"
        params.append(f"%{title}%")
    if min_score:
        query += " AND h.similarity_score >= %s"
        params.append(min_score)
    if max_score:
        query += " AND h.similarity_score <= %s"
        params.append(max_score)

    query += " ORDER BY h.similarity_score DESC"

    cur.execute(query, tuple(params))
    documents = cur.fetchall()

    cur.close()
    conn.close()

    return render_template(
    'user/documents_detail.html',
    documents=documents,
    active_page='history'   # ✅ Tambahin ini
)



INDONESIAN_STOPWORDS = set([
    'diperlukan', 'hendaknya', 'tapi', 'dimungkinkan', 'hendaklah', 'umumnya', 'tambahnya', 'usai', 'katakan', 'sebagaimana', 'sekali', 'persoalan', 'waduh', 'bermaksud', 'jelaslah', 'ditanyai', 'tiba', 'terdahulu', 'menghendaki', 'tidak', 'sangatlah', 'kalaulah', 'rata', 'tadi', 'sendirinya', 'tersampaikan', 'sekadar', 'mengakhiri', 'mempergunakan', 'sedikit', 'sekali-kali', 'katakanlah', 'karenanya', 'oleh', 'semampunya', 'diakhirinya', 'kapanpun', 'setidaknya', 'disini', 'menaiki', 'tentunya', 'terbanyak', 'tak', 'secara', 'diibaratkannya', 'mengatakan', 'hendak', 'dikarenakan', 'sekarang', 'berturut', 'ditanyakan', 'terlihat', 'diperlukannya', 'sebuah', 'cuma', 'ingat-ingat', 'sesegera', 'mengerjakan', 'keinginan', 'berlebihan', 'apalagi', 'siapapun', 'enggaknya', 'lagi', 'diungkapkan', 'bisa', 'tentu', 'bersiap', 'dia', 'ia', 'ini', 'dituturkan', 'mendatang', 'semacam', 'sebenarnya', 'terutama', 'diibaratkan', 'tunjuk', 'inilah', 'diri', 'seterusnya', 'menandaskan', 'kenapa', 'dimulailah', 'mengibaratkan', 'wong', 'disinilah', 'bahkan', 'kelihatan', 'sudahkah', 'mempertanyakan', 'dalam', 'luar', 'memulai', 'mengucapkan', 'selalu', 'waktu', 'ataukah', 'wahai', 'beberapa', 'semuanya', 'mampu', 'sebagainya', 'memungkinkan', 'bukannya', 'jadi', 'menanyakan', 'percuma', 'bolehkah', 'sekurang-kurangnya', 'yakin', 'memperbuat', 'jadinya', 'belumlah', 'terdiri', 'menjadi', 'sekalipun', 'merekalah', 'melihat', 'terakhir', 'hari', 'wah', 'sesuatu', 'sebelum', 'mendapat', 'berapa', 'dulu', 'sudah', 'tidaklah', 'kurang', 'makanya', 'ditunjuk', 'akhiri', 'bila', 'sayalah', 'buat', 'segalanya', 'berjumlah', 'perlunya', 'apatah', 'begitukah', 'itu', 'cara', 'antara', 'sampaikan', 'amat', 'mulailah', 'tertentu', 'setibanya', 'tiga', 'maka', 'semasih', 'nyaris', 'masalah', 'sebaik-baiknya', 'pasti', 'tiba-tiba', 'awal', 'bermula', 'tegasnya', 'bukanlah', 'selamanya', 'bermacam', 'satu', 'merupakan', 'disampaikan', 'sebanyak', 'menuturkan', 'segera', 'diucapkan', 'mendatangi', 'dipergunakan', 'bertanya-tanya', 'berkata', 'memintakan', 'jelas', 'kapan', 'tanyanya', 'tetapi', 'anda', 'benar', 'semula', 'sejenak', 'perlu', 'semakin', 'memang', 'begini', 'kemudian', 'serupa', 'disebutkan', 'pun', 'turut', 'bahwasanya', 'pastilah', 'nanti', 'didatangkan', 'dan', 'sedangkan', 'dikira', 'tentang', 'tersebutlah', 'diminta', 'dituturkannya', 'cukup', 'lanjutnya', 'dibuatnya', 'ucapnya', 'baru', 'haruslah', 'meminta', 'dijelaskan', 'kelihatannya', 'lainnya', 'ada', 'ibaratnya', 'ingin', 'menyangkut', 'mendapatkan', 'pentingnya', 'dirinya', 'dialah', 'diantaranya', 'terjadilah', 'ditujukan', 'bahwa', 'nah', 'mengibaratkannya', 'terhadap', 'saat', 'ditanya', 'ikut', 'mulanya', 'bakalan', 'setiba', 'tiap', 'bagaimana', 'sela', 'diberikannya', 'hanya', 'mengingat', 'meski', 'sebutlah', 'diinginkan', 'kata', 'hingga', 'usah', 'dikatakannya', 'apabila', 'per', 'manakala', 'untuk', 'sebegini', 'yakni', 'bertanya', 'olehnya', 'dipersoalkan', 'digunakan', 'ibu', 'teringat-ingat', 'adalah', 'berikan', 'sedemikian', 'sepihak', 'tandasnya', 'tegas', 'berlainan', 'bekerja', 'dini', 'inikah', 'mendatangkan', 'seringnya', 'terjadi', 'belakang', 'lalu', 'bawah', 'kedua', 'berada', 'jelaskan', 'bersiap-siap', 'awalnya', 'asal', 'daripada', 'mungkinkah', 'boleh', 'tutur', 'tengah', 'kasus', 'berikutnya', 'masing-masing', 'keadaan', 'terjadinya', 'meyakini', 'juga', 'ditunjuki', 'manalagi', 'menunjukkan', 'namun', 'bertutur', 'sehingga', 'terus', 'jadilah', 'ternyata', 'sama-sama', 'ditandaskan', 'ibaratkan', 'mirip', 'melihatnya', 'berkali-kali', 'ataupun', 'nyatanya', 'dimulai', 'bagi', 'jawabnya', 'teringat', 'aku', 'tambah', 'sudahlah', 'inginkah', 'seluruh', 'terasa', 'berakhirlah', 'dipertanyakan', 'kan', 'menyampaikan', 'saling', 'dimisalkan', 'sementara', 'beginikah', 'memastikan', 'walaupun', 'dibuat', 'kitalah', 'berkehendak', 'bilakah', 'ujar', 'pertanyakan', 'sendiri', 'jauh', 'dipunyai', 'tanpa', 'kamu', 'menyebutkan', 'berkeinginan', 'seseorang', 'pernah', 'beri', 'siapa', 'termasuk', 'pantas', 'pertama-tama', 'kelamaan', 'memperkirakan', 'semasa', 'didapat', 'belakangan', 'malahan', 'misal', 'mengungkapkan', 'yang', 'menunjuknya', 'setinggi', 'jika', 'sekalian', 'sepantasnyalah', 'kecil', 'masa', 'mau', 'bolehlah', 'lebih', 'lewat', 'betulkah', 'menanti', 'dimaksudnya', 'sebelumnya', 'jumlahnya', 'ditegaskan', 'bukan', 'di', 'mempersiapkan', 'sebesar', 'sekecil', 'bagaimanapun', 'sedikitnya', 'melalui', 'lamanya', 'benarlah', 'misalkan', 'kapankah', 'tetap', 'lagian', 'andalah', 'mengenai', 'mulai', 'mereka', 'bersama-sama', 'selama', 'ucap', 'soal', 'banyak', 'berawal', 'misalnya', 'nantinya', 'berdatangan', 'diketahui', 'jangan', 'suatu', 'biasa', 'seluruhnya', 'menantikan', 'atau', 'diberi', 'seingat', 'adapun', 'diantara', 'sering', 'ditambahkan', 'tuturnya', 'jikalau', 'berlalu', 'sebaliknya', 'begitupun', 'naik', 'diucapkannya', 'kelima', 'sepanjang', 'setiap', 'toh', 'itulah', 'sebaiknya', 'rasa', 'akhir', 'bagaikan', 'panjang', 'bagai', 'lanjut', 'benarkah', 'macam', 'sejumlah', 'menanya', 'semisalnya', 'serta', 'berujar', 'dekat', 'amatlah', 'artinya', 'bagaimanakah', 'khususnya', 'bersama', 'tandas', 'sebisanya', 'sejauh', 'sekitar', 'telah', 'balik', 'itukah', 'terlalu', 'dimaksudkan', 'sesekali', 'sebutnya', 'katanya', 'tidakkah', 'disebutkannya', 'sesudah', 'tampak', 'kalian', 'secukupnya', 'jawab', 'saya', 'masih', 'melakukan', 'pak', 'pula', 'dengan', 'menunjuk', 'sinilah', 'kembali', 'agaknya', 'antaranya', 'jelasnya', 'mengucapkannya', 'gunakan', 'diperkirakan', 'semua', 'tadinya', 'bermacam-macam', 'sebetulnya', 'jangankan', 'apaan', 'caranya', 'berapapun', 'demi', 'diperbuat', 'diperbuatnya', 'kira-kira', 'menginginkan', 'keterlaluan', 'tempat', 'bakal', 'menegaskan', 'tertuju', 'perlukah', 'sebaik', 'kita', 'agar', 'ketika', 'terkira', 'kalau', 'keseluruhannya', 'cukupkah', 'paling', 'seberapa', 'dua', 'selain', 'menyiapkan', 'setelah', 'justru', 'diingat', 'akulah', 'berkenaan', 'walau', 'lah', 'beginian', 'akhirnya', 'dikatakan', 'berapalah', 'soalnya', 'menurut', 'tanyakan', 'menjawab', 'seorang', 'ditunjukkannya', 'apakah', 'tersebut', 'makin', 'mengapa', 'sebagai', 'hanyalah', 'sebegitu', 'cukuplah', 'bukankah', 'sambil', 'dimaksudkannya', 'sesuatunya', 'selama-lamanya', 'sesama', 'hal', 'terdapat', 'apa', 'ialah', 'baik', 'belum', 'setidak-tidaknya', 'bulan', 'menambahkan', 'lama', 'masalahnya', 'mempersoalkan', 'melainkan', 'dahulu', 'berapakah', 'kiranya', 'demikian', 'lain', 'seperlunya', 'tentulah', 'meskipun', 'selaku', 'agak', 'diakhiri', 'saatnya', 'depan', 'dong', 'ungkapnya', 'guna', 'sedang', 'saja', 'kesampaian', 'berikut', 'memisalkan', 'penting', 'mengetahui', 'sekaligus', 'akankah', 'karena', 'pertanyaan', 'harus', 'kemungkinannya', 'semaunya', 'para', 'bung', 'keduanya', 'lima', 'jumlah', 'menyatakan', 'siap', 'kinilah', 'dipastikan', 'memerlukan', 'keluar', 'sama', 'sini', 'datang', 'selanjutnya', 'sajalah', 'sesaat', 'diingatkan', 'dimulainya', 'kalaupun', 'mengingatkan', 'harusnya', 'setempat', 'diperlihatkan', 'inginkan', 'mengatakannya', 'menjelaskan', 'entahlah', 'merasa', 'kini', 'ke', 'pertama', 'seolah-olah', 'berbagai', 'terhadapnya', 'jawaban', 'dapat', 'kebetulan', 'sesudahnya', 'berturut-turut', 'sangat', 'sampai', 'padanya', 'waktunya', 'menanyai', 'demikianlah', 'biasanya', 'betul', 'disebut', 'dilakukan', 'kemungkinan', 'pada', 'padahal', 'empat', 'beginilah', 'sempat', 'minta', 'menuju', 'ditunjuknya', 'se', 'sekurangnya', 'sekadarnya', 'pihak', 'dilihat', 'seolah', 'seperti', 'kepadanya', 'dimaksud', 'bagian', 'enggak', 'punya', 'keseluruhan', 'mampukah', 'adanya', 'tepat', 'menanti-nanti', 'begitulah', 'terlebih', 'maupun', 'sewaktu', 'rasanya', 'semata', 'menunjuki', 'dari', 'kamulah', 'sejak', 'kala', 'sekitarnya', 'begitu', 'seharusnya', 'kok', 'sampai-sampai', 'ditunjukkan', 'mengira', 'masing', 'supaya', 'diketahuinya', 'pukul', 'menyeluruh', 'semata-mata', 'berlangsung', 'tahun', 'diberikan', 'rupanya', 'tampaknya', 'menggunakan', 'atas', 'bisakah', 'tinggi', 'kamilah', 'mempunyai', 'pihaknya', 'berakhir', 'sepertinya', 'ujarnya', 'dikerjakan', 'sana', 'ungkap', 'berakhirnya', 'seketika', 'siapakah', 'umum', 'meyakinkan', 'sebabnya', 'membuat', 'dijelaskannya', 'kira', 'kepada', 'yaitu', 'seenaknya', 'malah', 'ibarat', 'janganlah', 'memihak', 'memberi', 'berarti', 'semampu', 'entah', 'sebut', 'segala', 'mungkin', 'memperlihatkan', 'sekiranya', 'hampir', 'tanya', 'berupa', 'sebagian', 'akan', 'semisal', 'besar', 'sebab', 'sesampai', 'dijawab', 'ingat', 'asalkan', 'sepantasnya', 'setengah', 'tahu', 'antar', 'dilalui', 'mana', 'seusai', 'masihkah', 'mula', 'memberikan', 'sendirian', 'kami', 'dimintai', 'bapak'
        # Machine Learning Algorithms and Techniques
    'logistic regression', 'decision tree', 'random forest', 'gradient boosting', 'xgboost', 'lightgbm', 'catboost','knn', 'k-nearest neighbors', 'svm', 'support vector machine', 'naive bayes', 'k-means', 'pca', 'principal component analysis','linear regression', 'ridge regression', 'lasso regression', 'elastic net', 'bayesian optimization',
    'ensemble methods', 'overfitting', 'underfitting', 'cross-validation', 'hyperparameter tuning', 'model evaluation',
    
    # Natural Language Processing (NLP) terms
    'tf-idf', 'word2vec', 'doc2vec', 'bert', 'roberta', 'distilbert', 'transformer', 'bert-based', 'rnn', 'lstm', 'gru','textblob', 'nltk', 'spacy', 'sastrawi', 'tokenization', 'lemmatization', 'stemming', 'stopword removal','n-grams', 'bag of words', 'word embeddings', 'sentence embeddings', 'dependency parsing', 'pos tagging', 'chunking',
    'entity recognition', 'named entity recognition', 'text classification', 'sentiment analysis', 'topic modeling','text summarization', 'machine translation', 'word sense disambiguation',
    
    # Computer Vision terms
    'cnn', 'convolutional neural network', 'resnet', 'efficientnet', 'mobilenet', 'vgg', 'alexnet', 'yolo', 'faster rcnn','object detection', 'semantic segmentation', 'instance segmentation', 'image classification', 'image recognition',
    'feature extraction', 'opencv', 'image processing', 'image augmentation', 'mask rcnn', 'style transfer', 'autoencoder','generative adversarial networks', 'gan', 'deep neural network', 'artificial neural network', 'r-cnn', 'detector', 'image generation', 'super resolution', 'semantic segmentation', 'deepfake', 'scene parsing',
    
    # Deep Learning Terms
    'deep learning', 'artificial neural network', 'autoencoder', 'generative adversarial networks', 'reinforcement learning',
    'backpropagation', 'gradient descent', 'activation function', 'loss function', 'batch normalization', 'dropout',
    'convolutional layer', 'pooling layer', 'fully connected layer', 'softmax', 'relu', 'sigmoid', 'tanh', 'adam optimizer',
    'vanishing gradient', 'exploding gradient', 'epoch', 'batch size', 'learning rate', 'momentum', 'adam',
    'rmsprop', 'sgd', 'optimizers', 'training set', 'validation set', 'test set', 'weights', 'bias', 'forward pass',
    'backpropagation', 'gradient updates', 'overfitting', 'underfitting', 'regularization', 'fine-tuning', 'transfer learning',
    
    # Frameworks and Libraries
    'tensorflow', 'keras', 'pytorch', 'theano', 'mxnet', 'caffe', 'tensorflow lite', 'tflite', 'scikit-learn', 'scipy',
    'pandas', 'numpy', 'matplotlib', 'seaborn', 'xgboost', 'lightgbm', 'catboost', 'sklearn', 'openCV', 'spaCy', 'gensim',
    'textblob', 'fastai', 'huggingface', 'pillow', 'cv2', 'nltk', 'chainer', 'cudnn', 'cupy', 'numba', 'plotly', 'dash',
    
    # Data Science and AI Concepts
    'machine learning', 'deep learning', 'artificial intelligence', 'natural language processing', 'nlp', 'data science',
    'data engineering', 'big data', 'data mining', 'predictive modeling', 'model training', 'feature engineering',
    'feature selection', 'clustering', 'regression', 'classification', 'time series', 'predictive analytics', 'analytics',
    'model deployment', 'model interpretability', 'model explainability', 'data preprocessing', 'data wrangling', 'data cleaning',
    'data transformation', 'data augmentation', 'bias-variance tradeoff', 'support vector machine', 'dimensionality reduction',
    'turing test', 'reinforcement learning', 'unsupervised learning', 'supervised learning', 'semi-supervised learning',
    'k-fold cross-validation', 'precision', 'recall', 'f1-score', 'accuracy', 'confusion matrix', 'roc curve', 'pr curve',
    'hyperparameter tuning', 'cross-validation', 'ensemble methods', 'stacking', 'boosting', 'bagging', 'hyperparameter optimization',
    'model selection', 'hyperparameter search', 'grid search', 'random search', 'learning curve',
    
    # AI Models
    'bert', 'gpt', 't5', 'roberta', 'electra', 'xlmr', 'bert-base', 'bert-large', 'distilbert', 'albert', 'gpt-2', 'gpt-3',
    'xlm-roberta', 'turing-nlg', 'gpt-neo', 't5-small', 't5-base', 't5-large', 'bertweet', 'sentencebert', 'fasttext', 'elmo',
    
    # Statistical Methods
    'bayesian statistics', 'hypothesis testing', 'probability distribution', 'normal distribution', 'poisson distribution','binomial distribution', 'confidence interval', 'hypothesis test', 'p-value', 'anova', 't-test', 'chi-squared test','regression analysis', 'multivariate analysis', 'p-value', 'covariance', 'correlation', 'statistical significance',    'linear regression', 'logistic regression', 'bayesian regression', 'gaussian distribution', 'z-score', 'multivariate analysis','algoritma', 'model', 'metode', 'proses', 'data', 'sistem', 'pemrograman', 'jaringan', 'komputer','framework', 'machine', 'learning', 'deep', 'learning', 'neural', 'network', 'klasifikasi', 'regresi', 
    'optimasi', 'teknologi', 'input', 'output', 'analisis', 'graf', 'fitur', 'transformasi', 'pengenalan', 'data', 'training', 'testing', 'validasi', 'akurasi', 'precision', 'recall', 'f1', 'score', 'evaluasi',     'cross', 'validation', 'pembelajaran', 'supervised', 'unsupervised', 'reinforcement', 'tuning', 'hyperparameter','kinerja', 'performansi', 'neural', 'network', 'training', 'epoch', 'batch', 'size', 'learning', 'rate', 
    'backpropagation', 'hidden', 'layer', 'convolutional', 'residual', 'dropout', 'activation', 'function',     'autoencoder', 'cnn', 'rnn', 'lstm', 'gru', 'decision', 'tree', 'svm', 'kmeans', 'knn', 'kfold', 'gradient', 'boosting', 'xgboost', 'lightgbm', 'catboost', 'naive', 'bayes', 'random', 'forest', 'clustering',     'regressor', 'classifier', 'loss', 'mean', 'squared', 'error', 'logistic', 'regression', 'support', 'vector', 'machine', 'feature', 'scaling', 'normalization', 'standardization', 'matrix', 'tensor', 'vector',     'classification', 'token', 'vectorization', 'embedding', 'word2vec', 'doc2vec', 'tfidf', 'bag', 'of', 'words', 
    'nlp', 'tokenization', 'lemmatization', 'stemming', 'stopword', 'phrase', 'syntax', 'semantic', 'context', 'language', 'models', 'text', 'analysis', 'pattern', 'recognition', 'computational', 'geometry',     'cloud', 'computing', 'iot', 'internet', 'things', 'ai', 'artificial', 'intelligence', 'big', 'data', 'analysis', 
    'hadoop', 'spark', 'dataframe', 'sql', 'nosql', 'query', 'database', 'mongodb', 'mysql', 'postgresql',     'cassandra', 'elasticsearch', 'index', 'search', 'algorithm', 'optimization', 'heuristic', 'dijkstra', 'bellman', 'ford', 'a_star', 'graph', 'search', 'dynamical', 'systems', 'complexity', 'theory', 'entropi', 
    'compression', 'binarization', 'parallel', 'distributed', 'computation', 'gpu', 'cpu', 'hardware', 'architecture', 'cloud', 'server', 'client', 'api', 'endpoint', 'microservices', 'containers', 'docker', 'kubernetes',     'restful', 'json', 'xml', 'rpc', 'async', 'sync', 'queue', 'queueing', 'job', 'distributed', 'system', 
    'event', 'driven', 'message', 'broker', 'mqtt', 'web', 'server', 'load', 'balancer', 'client', 'side', 'cloud',     'function', 'serverless', 'streaming', 'etl', 'pipeline', 'database', 'transaction', 'scalability', 
    'performance', 'latency', 'response', 'throughput', 'continuous', 'deployment', 'monitoring', 'logging', 'metrics', 'alert', 'security', 'encryption', 'authentication', 'authorization', 'firewall',     'proxy', 'vpn', 'malware', 'hacking', 'phishing', 'ddos', 'cryptocurrency', 'blockchain', 'distributed', 
    'ledger', 'smart', 'contract', 'hash', 'public', 'private', 'key', 'bitcoin', 'ethereum', 'crypto',     'network', 'protocol', 'tcp', 'ip', 'http', 'https', 'rest', 'soap', 'smtp', 'udp', 'ssl', 'tls', 'ipsec', 
    'vpn', 'nft', 'defi', 'decentralized', 'token', 'blockchain', 'tokens', 'miners', 'hashrate', 'node', 
   'websocket', 'mqtt', 'rpc', 'rest', 'graphql', 'json', 'graph', 'api', 'microservices', 'container',     'docker', 'kubernetes', 'cloud', 'cloud', 'aws', 'azure', 'gcp', 'google', 'compute', 'virtualization',    'containers', 'terraform', 'jenkins', 'ci', 'cd', 'azure', 'github', 'gitlab', 'devops', 'automated', 
    'deployment', 'integration', 'testing', 'monitoring', 'logging', 'stack', 'kafka', 'rabbitmq', 'event',     'streaming', 'data', 'architecture', 'service', 'architecture', 'data', 'cloud', 'serverless', 'api', 'integrations', 'restful', 'python', 'flask', 'django', 'javascript', 'nodejs', 'typescript', 'html', 
    'css', 'angular', 'react', 'vue', 'frontend', 'backend', 'fullstack', 'database', 'redis', 'mongodb',     'mysql', 'postgresql', 'sqlite', 'nosql', 'datastore', 'table', 'row', 'column', 'field', 'key', 'value',
    'query', 'index', 'schema', 'relation', 'link', 'metadata', 'authorization', 'deletion', 'foreign', 'primary',
    'constraint', 'transaction', 'join', 'sqlalchemy', 'orm', 'mysql', 'joins', 'groupby', 'orderby', 'combine',    'split', 'merge', 'indexing', 'group', 'cluster', 'inference', 'batch', 'dataloader', 'tensor', 'model',
    'pytorch', 'tensorflow', 'keras', 'scikit-learn', 'opencv', 'computervision', 'image', 'classifier',    'segmentation', 'object', 'detection', 'opencv', 'imageprocessing', 'opencv', 'transform', 'reshape',    'gradient', 'convolution', 'object', 'modeling', 'kernel', 'multi', 'multihead', 'attention', 'decoder', 
    'encoder', 'rl', 'agent', 'reinforcement', 'decay', 'exploration', 'exploitation', 'learning', 'scheduler',     'update', 'reward', 'epsilon', 'bandit', 'matplotlib', 'seaborn', 'plotly', 'numpy', 'pandas', 'array', 
    'linear', 'regression', 'random', 'forest', 'xgboost', 'lightgbm', 'svm', 'classification', 'crossvalidation',    'validation', 'dataset', 'time', 'series', 'anomaly', 'outlier', 'network', 'anomaly', 'model', 'fitting', 
    'outlier', 'detection', 'normalization', 'transformation', 'zscore', 'mean', 'median', 'sklearn', 'regression',
    'algoritma', 'model', 'metode', 'proses', 'data', 'sistem', 'pemrograman', 'jaringan', 'komputer',     'framework', 'machine', 'learning', 'deep', 'learning', 'neural', 'network', 'klasifikasi', 'regresi', 
    'optimasi', 'teknologi', 'input', 'output', 'analisis', 'graf', 'fitur', 'transformasi', 'pengenalan',     'data', 'training', 'testing', 'validasi', 'akurasi', 'precision', 'recall', 'f1', 'score', 'evaluasi', 
    'cross', 'validation', 'pembelajaran', 'supervised', 'unsupervised', 'reinforcement', 'tuning', 'hyperparameter',
    'kinerja', 'performansi', 'neural', 'network', 'training', 'epoch', 'batch', 'size', 'learning', 'rate',     'backpropagation', 'hidden', 'layer', 'convolutional', 'residual', 'dropout', 'activation', 'function', 
    'autoencoder', 'cnn', 'rnn', 'lstm', 'gru', 'decision', 'tree', 'svm', 'kmeans', 'knn', 'kfold',     'gradient', 'boosting', 'xgboost', 'lightgbm', 'catboost', 'naive', 'bayes', 'random', 'forest', 'clustering', 
    'regressor', 'classifier', 'loss', 'mean', 'squared', 'error', 'logistic', 'regression', 'support', 'vector',     'machine', 'feature', 'scaling', 'normalization', 'standardization', 'matrix', 'tensor', 'vector', 
    'classification', 'token', 'vectorization', 'embedding', 'word2vec', 'doc2vec', 'tfidf', 'bag', 'of', 'words', 
    'nlp', 'tokenization', 'lemmatization', 'stemming', 'stopword', 'phrase', 'syntax', 'semantic',     'context', 'language', 'models', 'text', 'analysis', 'pattern', 'recognition', 'computational', 'geometry', 
    'cloud', 'computing', 'iot', 'internet', 'things', 'ai', 'artificial', 'intelligence', 'big', 'data', 'analysis', 
    'hadoop', 'spark', 'dataframe', 'sql', 'nosql', 'query', 'database', 'mongodb', 'mysql', 'postgresql',     'cassandra', 'elasticsearch', 'index', 'search', 'algorithm', 'optimization', 'heuristic', 'dijkstra', 
    'bellman', 'ford', 'a_star', 'graph', 'search', 'dynamical', 'systems', 'complexity', 'theory', 'entropi', 
    'compression', 'binarization', 'parallel', 'distributed', 'computation', 'gpu', 'cpu', 'hardware', 'architecture', 
    'cloud', 'server', 'client', 'api', 'endpoint', 'microservices', 'containers', 'docker', 'kubernetes',     'restful', 'json', 'xml', 'rpc', 'async', 'sync', 'queue', 'queueing', 'job', 'distributed', 'system', 
    'event', 'driven', 'message', 'broker', 'mqtt', 'web', 'server', 'load', 'balancer', 'client', 'side', 'cloud', 
    'function', 'serverless', 'streaming', 'etl', 'pipeline', 'database', 'transaction', 'scalability', 
    'performance', 'latency', 'response', 'throughput', 'continuous', 'deployment', 'monitoring',     'logging', 'metrics', 'alert', 'security', 'encryption', 'authentication', 'authorization', 'firewall', 
    'proxy', 'vpn', 'malware', 'hacking', 'phishing', 'ddos', 'cryptocurrency', 'blockchain', 'distributed', 
    'ledger', 'smart', 'contract', 'hash', 'public', 'private', 'key', 'bitcoin', 'ethereum', 'crypto',     'network', 'protocol', 'tcp', 'ip', 'http', 'https', 'rest', 'soap', 'smtp', 'udp', 'ssl', 'tls', 'ipsec', 
    'vpn', 'nft', 'defi', 'decentralized', 'token', 'blockchain', 'tokens', 'miners', 'hashrate', 'node', 
    'websocket', 'mqtt', 'rpc', 'rest', 'graphql', 'json', 'graph', 'api', 'microservices', 'container', 
    'docker', 'kubernetes', 'cloud', 'cloud', 'aws', 'azure', 'gcp', 'google', 'compute', 'virtualization',
    'containers', 'terraform', 'jenkins', 'ci', 'cd', 'azure', 'github', 'gitlab', 'devops', 'automated', 
    'deployment', 'integration', 'testing', 'monitoring', 'logging', 'stack', 'kafka', 'rabbitmq', 'event', 
    'streaming', 'data', 'architecture', 'service', 'architecture', 'data', 'cloud', 'serverless', 
    'api', 'integrations', 'restful', 'python', 'flask', 'django', 'javascript', 'nodejs', 'typescript', 'html', 
    'css', 'angular', 'react', 'vue', 'frontend', 'backend', 'fullstack', 'database', 'redis', 'mongodb', 
    'mysql', 'postgresql', 'sqlite', 'nosql', 'datastore', 'table', 'row', 'column', 'field', 'key', 'value',
    'query', 'index', 'schema', 'relation', 'link', 'metadata', 'authorization', 'deletion', 'foreign', 'primary',
    'constraint', 'transaction', 'join', 'sqlalchemy', 'orm', 'mysql', 'joins', 'groupby', 'orderby', 'combine',
    'split', 'merge', 'indexing', 'group', 'cluster', 'inference', 'batch', 'dataloader', 'tensor', 'model',    'pytorch', 'tensorflow', 'keras', 'scikit-learn', 'opencv', 'computervision', 'image', 'classifier',
    'segmentation', 'object', 'detection', 'opencv', 'imageprocessing', 'opencv', 'transform', 'reshape',    'gradient', 'convolution', 'object', 'modeling', 'kernel', 'multi', 'multihead', 'attention', 'decoder',     'encoder', 'rl', 'agent', 'reinforcement', 'decay', 'exploration', 'exploitation', 'learning', 'scheduler', 
    'update', 'reward', 'epsilon', 'bandit', 'matplotlib', 'seaborn', 'plotly', 'numpy', 'pandas', 'array',     'linear', 'regression', 'random', 'forest', 'xgboost', 'lightgbm', 'svm', 'classification', 'crossvalidation',
    'validation', 'dataset', 'time', 'series', 'anomaly', 'outlier', 'network', 'anomaly', 'model', 'fitting',     'outlier', 'detection', 'normalization', 'transformation', 'zscore', 'mean', 'median', 'sklearn', 'regression',
    'algorithm', 'machine', 'learning', 'deep', 'learning', 'data', 'model', 'training', 'testing', 'validation', 'accuracy',
    'loss', 'function', 'optimization', 'feature', 'selection', 'feature', 'extraction', 'cross-validation', 'confusion', 'matrix',    'overfitting', 'underfitting', 'bias', 'variance', 'gradient', 'descent', 'stochastic', 'training', 'backpropagation', 'classification',
    'regression', 'supervised', 'unsupervised', 'reinforcement', 'unsupervised', 'supervised', 'learning', 'support', 'vector',
    'svm', 'xgboost', 'catboost', 'lightgbm', 'kmeans', 'pca', 'tf-idf', 'vectorization', 'word2vec', 'doc2vec', 'bag', 'of', 'words',    'nlp', 'tokenization', 'lemmatization', 'stemming', 'stopword', 'embedding', 'word', 'context', 'rnn', 'lstm', 'gru', 'transformer',
    'bert', 'roberta', 'gpt', 'tensor', 'flow', 'keras', 'pytorch', 'neural', 'network', 'training', 'epochs', 'batch', 'size', 'model',
    'layer', 'activation', 'function', 'gpu', 'cpu', 'data', 'preprocessing', 'normalization', 'standardization', 'scikit-learn',    'random', 'forest', 'decision', 'tree', 'naive', 'bayes', 'support', 'vector', 'machine', 'svm', 'knn', 'k-nearest', 'neighbors',
    'logistic', 'regression', 'linear', 'regression', 'ensemble', 'methods', 'classification', 'accuracy', 'precision', 'recall', 'f1',
    'score', 'model', 'selection', 'hyperparameter', 'optimization', 'pipeline', 'batch', 'size', 'learning', 'rate', 'optimizer',
    'adam', 'sgd', 'rmsprop', 'metrics', 'roc', 'curve', 'pr', 'curve', 'precision', 'recall', 'f1', 'score', 'cross', 'entropy',
    'mse', 'mean', 'squared', 'error', 'mae', 'mean', 'absolute', 'error', 'time', 'series', 'clustering', 'dimensionality', 'reduction',
    'factorization', 'matrix', 'factorization', 'apriori', 'association', 'rules', 'dbscan', 'k-means', 'hierarchical', 'clustering',
    'ensemble', 'boosting', 'bagging', 'pipelining', 'batch', 'gradient', 'boosting', 'random', 'forest', 'neural', 'network',
    'data', 'visualization', 'matplotlib', 'seaborn', 'plotly', 't-sne', 'scatter', 'plot', 'line', 'plot', 'histogram', 'dataframe',
    'pandas', 'numpy', 'matrix', 'tensor', 'data', 'column', 'table', 'row', 'schema', 'join', 'query', 'database', 'sql', 'mongodb',
    'nosql', 'join', 'graph', 'mlflow', 'model', 'serving', 'training', 'classification', 'outliers', 'anomaly', 'detection', 'residual',
    'modeling', 'batch', 'job', 'cloud', 'hadoop', 'spark', 'airflow', 'etl', 'pipeline', 'dataset', 'batch', 'model', 'store', 'training',
    'deployment', 'cloud', 'serverless', 'api', 'data', 'streaming', 'batch', 'processing', 'cloud', 'azure', 'aws', 'gcp', 'terraform',
    'jenkins', 'ci', 'cd', 'docker', 'containers', 'rest', 'graphql', 'kubernetes', 'queue', 'asynchronous', 'queue', 'message', 'broker',
    'kafka', 'rabbitmq', 'mqtt', 'event', 'driven', 'architecture', 'microservices', 'cloud', 'devops', 'agile', 'kanban', 'scrum',
    'continuous', 'integration', 'continuous', 'deployment', 'integration', 'cloud', 'compute', 'virtualization', 'docker', 'k8s', 'nginx', 'apache', 'cloud', 'vps', 'host', 'web', 'server', 'api', 'access', 'restful', 'http', 'https', 'cookies', 'oauth',
    'jwt', 'session', 'authentication', 'firewall', 'proxy', 'vpn', 'waf', 'ddos', 'malware', 'hacking', 'cyber', 'security', 'phishing',
    'cryptocurrency', 'blockchain', 'crypto', 'bitcoin', 'ethereum', 'ledger', 'wallet', 'decentralized', 'miner', 'hashrate', 'hash',
    'node', 'ethereum', 'block', 'blockchain', 'hash', 'public', 'key', 'private', 'key', 'ethereum', 'bitcoin', 'token', 'decentralized',
    'smart', 'contract', 'distributed', 'ledger', 'blockchain', 'digital', 'currency', 'futures', 'tokens', 'nft', 'defi', 'smart',
    'contracts', 'trading', 'stability', 'mempool', 'sharding', 'proof', 'work', 'stake', 'proof', 'consensus', 'miner', 'block',    'blockchain', 'icp', 'fiat', 'smart', 'tokens', 'airdrops', 'nft', 'data', 'storage', 'oracle', 'yarn', 'spark', 'hive', 'hadoop',
    'apache', 'airflow', 'data', 'cloud', 'architecture', 'containerization', 'web', 'database', 'data', 'design', 'pattern', 'models',
    'evolution', 'models', 'clustering', 'training', 'fitting', 'outliers', 'error', 'balancing', 'approach', 'computation', 'distributed',
    'event', 'driven', 'api', 'services', 'cloud', 'cloud', 'internet', 'things', 'iot', 'edge', 'computing', 'devices', 'sensors',     'sensor', 'edge', 'computing', 'network', 'latency', 'response', 'real-time', 'sensor', 'communication', 'iot', 'iotcloud', 'datastore'     'algorithm', 'machine', 'learning', 'deep', 'learning', 'data', 'model', 'training', 'testing', 'validation', 'accuracy',
    'loss', 'function', 'optimization', 'feature', 'selection', 'feature', 'extraction', 'cross-validation', 'confusion', 'matrix',
    'overfitting', 'underfitting', 'bias', 'variance', 'gradient', 'descent', 'stochastic', 'training', 'backpropagation', 'classification',    'regression', 'supervised', 'unsupervised', 'reinforcement', 'unsupervised', 'supervised', 'learning', 'support', 'vector',
    'svm', 'xgboost', 'catboost', 'lightgbm', 'kmeans', 'pca', 'tf-idf', 'vectorization', 'word2vec', 'doc2vec', 'bag', 'of', 'words',
    'nlp', 'tokenization', 'lemmatization', 'stemming', 'stopword', 'embedding', 'word', 'context', 'rnn', 'lstm', 'gru', 'transformer',
    'bert', 'roberta', 'gpt', 'tensor', 'flow', 'keras', 'pytorch', 'neural', 'network', 'training', 'epochs', 'batch', 'size', 'model',
    'layer', 'activation', 'function', 'gpu', 'cpu', 'data', 'preprocessing', 'normalization', 'standardization', 'scikit-learn',    'random', 'forest', 'decision', 'tree', 'naive', 'bayes', 'support', 'vector', 'machine', 'svm', 'knn', 'k-nearest', 'neighbors',
    'logistic', 'regression', 'linear', 'regression', 'ensemble', 'methods', 'classification', 'accuracy', 'precision', 'recall', 'f1',    'score', 'model', 'selection', 'hyperparameter', 'optimization', 'pipeline', 'batch', 'size', 'learning', 'rate', 'optimizer',    'adam', 'sgd', 'rmsprop', 'metrics', 'roc', 'curve', 'pr', 'curve', 'precision', 'recall', 'f1', 'score', 'cross', 'entropy',
    'mse', 'mean', 'squared', 'error', 'mae', 'mean', 'absolute', 'error', 'time', 'series', 'clustering', 'dimensionality', 'reduction',    'factorization', 'matrix', 'factorization', 'apriori', 'association', 'rules', 'dbscan', 'k-means', 'hierarchical', 'clustering',    'ensemble', 'boosting', 'bagging', 'pipelining', 'batch', 'gradient', 'boosting', 'random', 'forest', 'neural', 'network',
    'data', 'visualization', 'matplotlib', 'seaborn', 'plotly', 't-sne', 'scatter', 'plot', 'line', 'plot', 'histogram', 'dataframe',    'pandas', 'numpy', 'matrix', 'tensor', 'data', 'column', 'table', 'row', 'schema', 'join', 'query', 'database', 'sql', 'mongodb',    'nosql', 'join', 'graph', 'mlflow', 'model', 'serving', 'training', 'classification', 'outliers', 'anomaly', 'detection', 'residual',
    'modeling', 'batch', 'job', 'cloud', 'hadoop', 'spark', 'airflow', 'etl', 'pipeline', 'dataset', 'batch', 'model', 'store', 'training',    'deployment', 'cloud', 'serverless', 'api', 'data', 'streaming', 'batch', 'processing', 'cloud', 'azure', 'aws', 'gcp', 'terraform','jenkins', 'ci', 'cd', 'docker', 'containers', 'rest', 'graphql', 'kubernetes', 'queue', 'asynchronous', 'queue', 'message', 'broker',
    'kafka', 'rabbitmq', 'mqtt', 'event', 'driven', 'architecture', 'microservices', 'cloud', 'devops', 'agile', 'kanban', 'scrum',    'continuous', 'integration', 'continuous', 'deployment', 'integration', 'cloud', 'compute', 'virtualization', 'docker', 'k8s',    'nginx', 'nginx', 'apache', 'cloud', 'vps', 'host', 'web', 'server', 'api', 'access', 'restful', 'http', 'https', 'cookies', 'oauth',    'jwt', 'session', 'authentication', 'firewall', 'proxy', 'vpn', 'waf', 'ddos', 'malware', 'hacking', 'cyber', 'security', 'phishing',    'cryptocurrency', 'blockchain', 'crypto', 'bitcoin', 'ethereum', 'ledger', 'wallet', 'decentralized', 'miner', 'hashrate', 'hash',
    'node', 'ethereum', 'block', 'blockchain', 'hash', 'public', 'key', 'private', 'key', 'ethereum', 'bitcoin', 'token', 'decentralized',    'smart', 'contract', 'distributed', 'ledger', 'blockchain', 'digital', 'currency', 'futures', 'tokens', 'nft', 'defi', 'smart',    'contracts', 'trading', 'stability', 'mempool', 'sharding', 'proof', 'work', 'stake', 'proof', 'consensus', 'miner', 'block',    'blockchain', 'icp', 'fiat', 'smart', 'tokens', 'airdrops', 'nft', 'data', 'storage', 'oracle', 'yarn', 'spark', 'hive', 'hadoop',    'apache', 'airflow', 'data', 'cloud', 'architecture', 'containerization', 'web', 'database', 'data', 'design', 'pattern', 'models',    'evolution', 'models', 'clustering', 'training', 'fitting', 'outliers', 'error', 'balancing', 'approach', 'computation', 'distributed',    'event', 'driven', 'api', 'services', 'cloud', 'cloud', 'internet', 'things', 'iot', 'edge', 'computing', 'devices', 'sensors',     'sensor', 'edge', 'computing', 'network', 'latency', 'response', 'real-time', 'sensor', 'communication', 'iot', 'iotcloud', 'datastore'
])



def mark_similar_words_in_original(original_text, similar_tokens):
    """
    Menandai kata-kata dalam teks asli berdasarkan token yang dianggap similar.

    Args:
        original_text (str): Teks asli dari dokumen.
        similar_tokens (set): Token-token mirip dari dokumen lain.

    Returns:
        list: Daftar tuple (kata, is_similar)
    """
    from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
    import re

    # Inisialisasi stemmer
    factory = StemmerFactory()
    stemmer = factory.create_stemmer()

    # Split kata: kata, spasi, atau tanda baca
    words = re.findall(r'\b\w+\b|\s+|[^\w\s]', original_text)

    marked_words = []

    for word in words:
        # Deteksi kata (bukan spasi atau tanda baca)
        if re.match(r'\b\w+\b', word):
            word_clean = word.strip().lower()

            # Abaikan token pendek/terlalu umum
            if len(word_clean) <= 2:
                marked_words.append((word, False))
                continue

            # Lakukan stemming
            stemmed = stemmer.stem(word_clean)

            # Tandai jika ada dalam similar_tokens
            if stemmed in similar_tokens:
                marked_words.append((word, True))
            else:
                marked_words.append((word, False))
        else:
            # Pertahankan spasi dan tanda baca
            marked_words.append((word, False))

    return marked_words



# Contoh penggunaan fungsi dalam kode utama
def test_similarity(model, doc_id):
    try:
        doc_vector = model.dv[str(doc_id)]
        return model.dv.most_similar([doc_vector], topn=5)
    except KeyError:
        return f"Dokumen dengan ID {doc_id} tidak ditemukan dalam model."

    
# Function untuk menghasilkan PDF dengan highlight

import html  # tambahkan ini di bagian import atas
def create_highlighted_pdf(original_text, similar_tokens, output_path, user_id=None, filename_hint=None):
    """
    Membuat PDF baru dengan highlight + simpan hasil preprocessing ke JSON.

    Args:
        original_text (str): Teks asli dari dokumen
        similar_tokens (set): Token yang dianggap mirip
        output_path (str): Path PDF yang akan disimpan
        user_id (int, optional): Untuk membuat folder khusus user
        filename_hint (str, optional): Nama file JSON agar mudah ditelusuri
    """

    # ✅ Simpan hasil processing ke backend
    processed_dir = os.path.join('processed_texts')
    os.makedirs(processed_dir, exist_ok=True)

    json_filename = f"{filename_hint or 'result'}_highlight.json"
    json_path = os.path.join(processed_dir, json_filename)

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump({
            "original_text": original_text,
            "similar_tokens": list(similar_tokens),
            "output_path": output_path,
            "user_id": user_id
        }, f, ensure_ascii=False, indent=2)

    # 🟨 Tandai kata-kata mirip
    marked_words = mark_similar_words_in_original(original_text, similar_tokens)

    # ⬇️ PDF rendering
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=72
    )
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]

    content = []
    current_paragraph = ""

    for word, is_similar in marked_words:
        escaped_word = html.escape(word)
        if is_similar:
            current_paragraph += f'<font backcolor="yellow">{escaped_word}</font>'
        else:
            current_paragraph += escaped_word

        if word in ('\n', '\r\n'):
            if current_paragraph:
                p = Paragraph(current_paragraph, normal_style)
                content.append(p)
                content.append(Spacer(1, 12))
                current_paragraph = ""

    if current_paragraph:
        p = Paragraph(current_paragraph, normal_style)
        content.append(p)

    doc.build(content)
    return output_path
# Modifikasi function preprocess_text dan calculate_similarity untuk menyimpan token similar


def generate_highlighted_exum_pdf(user_id, exum_text, db_documents):
    similarities = calculate_similarity(exum_text, db_documents)


    # Get top 5 similar documents
    top_5_docs = similarities[:5]
    similarity_dicts = []

    # Collecting the tokens for each of the top 5 similar documents
    for doc_id, title, score, similar_tokens in top_5_docs:
        similarity_dicts.append({
            "doc_id": doc_id,
            "tokens": set(similar_tokens)  # Assuming similar_tokens is a list of similar words
        })
    
    # Combine Exum text with the highlighted words from top 5 documents
    combined_text = exum_text  # Assuming exum_text is the full text of the Exum document
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'results', str(user_id), 'highlighted_exum.pdf')

    # Create the combined highlighted PDF
    return create_combined_exum_highlighted_pdf(combined_text, similarity_dicts, output_path)


@app.route('/admin/view-exum/<int:user_id>/<filename>')
@role_required(['admin'])  # optional: jika ingin batasi hanya admin
def admin_view_exum(user_id, filename):
    result_filename = f"{os.path.splitext(filename)[0]}_exum_combined.pdf"
    result_path = os.path.join(app.config['UPLOAD_FOLDER'], 'results', str(user_id), result_filename)
    
    if os.path.exists(result_path):
        return send_file(result_path, mimetype='application/pdf', as_attachment=False)
    
    abort(404, description="File hasil exum tidak ditemukan.")
    # Buat nama file PDF menggunakan ID history agar unik
    base_name = os.path.splitext(file_name)[0]
    result_file = os.path.join(result_dir, f"{base_name}_doc{history_id}.pdf")

    # Buat PDF dengan highlight
    create_highlighted_pdf(original_text, similar_tokens, result_file)

    # Perbarui history dengan path file hasil
    cur.execute(
        "UPDATE history SET result_file_path = %s WHERE id = %s",
        (result_file, history_id)
    )
    conn.commit()

    cur.close()
    conn.close()

    return result_file


def get_documents_from_db(document_type='skripsi'):
    conn = psycopg2.connect(
        "dbname=db_similarity user=postgres password=admin host=localhost"
    )
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    # Ambil hanya dokumen sesuai tipe (skripsi atau exum)
    cur.execute(
        "SELECT id, title, file_text, vector, file_name FROM documents WHERE document_type = %s",
        (document_type,)
    )
    documents = cur.fetchall()

    cur.close()
    conn.close()

    return [{
        'id': doc['id'],
        'title': doc['title'],
        'file_text': doc['file_text'],
        'vector': doc['vector'],
        'file_name': doc['file_name']
    } for doc in documents]

    


# Store based on Push Command:
# user_id, document_id = doc_id, uploaded_file_name = filename,
# '', similarity_score = score, '', '',
# result_file_path = skripsi_highlight_path 
# save_to_history(user_id, doc_id, filename, score, skripsi_highlight_path)
def save_to_history(user_id, filename, total_similarity, similarities):
    conn = psycopg2.connect(
        "dbname=db_similarity user=postgres password=admin host=localhost"
    )
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    history_ids = []
    for doc_id, title, score, exum, exum_combined, skripsi_highlight_path in similarities:
        matched_text_str = str('u')
        upload_batch_id = 0

        cur.execute(
            """
            INSERT INTO history (user_id, document_id, 
            uploaded_file_name, uploaded_file_text, similarity_score, 
            matched_text, result_file_path, upload_batch_id) 
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s) RETURNING id
            """,
            (user_id, doc_id, exum, exum_combined, score, matched_text_str, skripsi_highlight_path, upload_batch_id)
        )

        result = cur.fetchone()
        if result and 'id' in result:
            history_id = result['id']
            history_ids.append((
                history_id,
                doc_id,
                exum,
                exum_combined,
                score,
                matched_text_str,
                skripsi_highlight_path,
                upload_batch_id
            ))

    conn.commit()
    cur.close()
    conn.close()
    return history_ids

@app.route('/upload', methods=['GET', 'POST'])
@role_required(['user'])  # atau ['admin'] kalau mau untuk admin
def upload_page():
    if request.method == 'POST':
        file = request.files['file']
        compare_to = request.form.get('compare_to', 'skripsi')

        if file:
            filename = secure_filename(file.filename)
            save_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'admin_documents')
            os.makedirs(save_dir, exist_ok=True)
            file_path = os.path.join(save_dir, filename)
            file.save(file_path)

            return redirect(url_for('check_similarity_from_saved', filename=filename, compare_to=compare_to))

    return render_template('user/upload.html', active_page='upload')



    # hasil similarity dan highlight
    # Fungsi untuk menghasilkan teks HTML dari hasil similarity
def generate_similarity_output(similarities, db_docs):
        """
        Menghasilkan teks HTML untuk hasil similarity dan token-token untuk highlight.

        Args:
            similarities: List tuple (doc_id, score, matched_tokens)
            db_docs: List dokumen dari database

        Returns:
            output_text: HTML string hasil similarity
            highlighted_results: List (doc_id, matched_tokens)
        """
        output_text = ""
        highlighted_results = []

        for doc_id, score, matched_tokens in similarities:
            doc = next((d for d in db_docs if d['id'] == doc_id), None)
            if doc:
                matched_text = ', '.join(matched_tokens)
                output_text += f"<p><strong>{doc['title']}</strong> - Skor: {score:.2f}%<br>Kata mirip: {matched_text}</p>"
                highlighted_results.append((doc['id'], matched_tokens))

        return output_text, highlighted_results


@app.route('/check_similarity', methods=['POST'])
@role_required(['user'])
def check_similarity():
    if 'file' not in request.files or request.files['file'].filename == '':
        flash("No file selected", 'error')
        return redirect(request.url)

    file = request.files['file']
    compare_to = request.form.get('compare_to', 'skripsi')
    print("User membandingkan dengan:", compare_to)
    print("DEBUG SESSION di /check_similarity:", dict(session))

    filename = f"{int(time.time())}_{secure_filename(file.filename)}"
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    print("File uploaded:", filename)
    print("File saved to:", file_path)

    # Ekstrak teks dari file
    original_content = ""
    if filename.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as f:
            original_content = f.read()
    elif filename.endswith('.docx'):
        doc = Document(file_path)
        original_content = "\n".join([p.text for p in doc.paragraphs])
    elif filename.endswith('.pdf'):
        with pdfplumber.open(file_path) as pdf:
            original_content = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

    # Ambil dokumen pembanding dari DB
    db_docs = get_documents_from_db(compare_to)
    similarities = calculate_similarity(original_content, db_docs)
    total_similarity = sum([sim['combined_score'] for sim in similarities]) / len(similarities) if similarities else 0

    user_id = session.get('user_id')
    if not user_id:
        flash("User not logged in", 'error')
        return redirect('/login')

    # Ambil token mirip dari kalimat dengan dokumen pertama
    first_doc_text = db_docs[0]['file_text'] if db_docs else ""
    all_similar_tokens = get_stemmed_tokens_from_similar_sentences(original_content, first_doc_text, model)

    # Simpan hasil highlight ke JSON
    json_path = os.path.join('processed_texts', generate_highlight_filename(filename))

    os.makedirs('processed_texts', exist_ok=True)
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump({
            "original_text": original_content,
            "similar_tokens": list(all_similar_tokens),
            "filename": filename,
            "user_id": user_id
        }, f, ensure_ascii=False, indent=2)

    # Highlight dokumen yang diunggah (Exum/User)
    highlighted_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'results', str(user_id))
    os.makedirs(highlighted_dir, exist_ok=True)
    output_pdf_filename = filename.replace('.pdf', '_combined.pdf')
    output_pdf_path = os.path.join(highlighted_dir, output_pdf_filename)
    highlight_pdf_based_on_tokens(file_path, output_pdf_path, all_similar_tokens)

    result_files = []

    for sim in similarities:
        doc_id = sim['id']
        title = sim['title']
        score = sim['combined_score']
        similar_tokens = sim['similar_tokens']
        db_doc = next((doc for doc in db_docs if doc['id'] == doc_id), None)
        if not db_doc:
            continue

        clean_tokens = remove_html_tags(' '.join(similar_tokens))
        result_file = save_result_pdf(
            user_id=user_id,
            file_name=filename,
            original_text=db_doc['file_text'],
            similar_tokens=clean_tokens,
            history_id=doc_id
        )

        # Highlight dokumen skripsi dari database
        skripsi_highlight_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'skripsi_highlights', str(user_id))
        os.makedirs(skripsi_highlight_dir, exist_ok=True)

        skripsi_highlight_filename = generate_unique_filename()
        skripsi_highlight_path = os.path.join(skripsi_highlight_dir, skripsi_highlight_filename)

        file_name_db = db_doc['file_name']
        skripsi_input_path = os.path.join(app.config['UPLOAD_FOLDER'], 'admin_documents', file_name_db)

        if file_name_db and os.path.exists(skripsi_input_path):
            highlight_pdf_based_on_tokens(
                pdf_path=skripsi_input_path,
                output_path=skripsi_highlight_path,
                similar_tokens=similar_tokens
            )

        result_files.append((
            doc_id,
            db_doc['title'],
            score,
            filename,
            output_pdf_filename,
            skripsi_highlight_filename
        ))

    print("✅ Semua proses selesai")
    save_to_history(user_id, output_pdf_filename, total_similarity, result_files)

    
    return render_template(
        'similarity_results.html',
        similarities=result_files,
        total_similarity=total_similarity,
        highlighted_pdf=output_pdf_filename,
        active_page='upload'   
    )

@app.route('/check_similarity_from_saved')
@role_required(['user'])  # atau ['admin']
def check_similarity_from_saved():
    filename = request.args.get('filename')
    compare_to = request.args.get('compare_to', 'skripsi')
    user_id = session.get('user_id')

    if not filename or not user_id:
        flash("File atau sesi tidak valid", 'danger')
        return redirect('/upload')

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'admin_documents', filename)
    if not os.path.exists(file_path):
        flash("File tidak ditemukan", 'danger')
        return redirect('/upload')

    return process_similarity_from_path(file_path, compare_to, user_id)

def process_similarity_from_path(file_path, compare_to, user_id):
    filename = os.path.basename(file_path)

    # Ekstraksi teks
    if filename.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as f:
            original_content = f.read()
    elif filename.endswith('.docx'):
        doc = Document(file_path)
        original_content = "\n".join([p.text for p in doc.paragraphs])
    elif filename.endswith('.pdf'):
        with pdfplumber.open(file_path) as pdf:
            original_content = "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
    else:
        flash("Format file tidak didukung", 'danger')
        return redirect('/upload')

    db_docs = get_documents_from_db(compare_to)
    similarities = calculate_similarity(original_content, db_docs)
    total_similarity = sum([sim['combined_score'] for sim in similarities]) / len(similarities) if similarities else 0

    # Simpan token highlight
    processed_dir = os.path.join('processed_texts')
    os.makedirs(processed_dir, exist_ok=True)

    json_filename = generate_highlight_filename(filename)
    json_path = os.path.join(processed_dir, json_filename)

    all_similar_tokens = set()
    for sim in similarities:
        all_similar_tokens.update(sim['similar_tokens'])

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump({
            "original_text": original_content,
            "similar_tokens": list(all_similar_tokens),
            "filename": filename,
            "user_id": user_id
        }, f, ensure_ascii=False, indent=2)

    # Highlight hasil exum
    highlighted_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'results', str(user_id))
    os.makedirs(highlighted_dir, exist_ok=True)
    output_pdf_filename = filename.replace('.pdf', '_combined.pdf')
    output_pdf_path = os.path.join(highlighted_dir, output_pdf_filename)
    create_highlighted_pdf(original_content, all_similar_tokens, output_pdf_path)

    # Proses dokumen DB
    result_files = []
    for sim in similarities:
        doc_id = sim['id']
        title = sim['title']
        score = sim['combined_score']
        similar_tokens = sim['similar_tokens']

        db_doc = next((doc for doc in db_docs if doc['id'] == doc_id), None)
        if not db_doc:
            continue

        clean_tokens = remove_html_tags(' '.join(similar_tokens))
        result_file = save_result_pdf(
            user_id=user_id,
            file_name=filename,
            original_text=db_doc['file_text'],
            similar_tokens=clean_tokens,
            history_id=doc_id
        )

        # Highlight skripsi dari DB
        skripsi_highlight_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'skripsi_highlights', str(user_id))
        os.makedirs(skripsi_highlight_dir, exist_ok=True)

        skripsi_highlight_filename = generate_unique_filename()
        skripsi_highlight_path = os.path.join(skripsi_highlight_dir, skripsi_highlight_filename)

        file_name_db = db_doc['file_name']
        skripsi_input_path = os.path.join(app.config['UPLOAD_FOLDER'], 'admin_documents', file_name_db)

        if file_name_db and os.path.exists(skripsi_input_path):
            highlight_pdf_based_on_tokens(
                pdf_path=skripsi_input_path,
                output_path=skripsi_highlight_path,
                similar_tokens=similar_tokens
            )

        result_files.append((
            doc_id,
            title,
            score,
            filename,
            output_pdf_filename,
            skripsi_highlight_filename
        ))

    save_to_history(user_id, output_pdf_filename, total_similarity, result_files)

    return render_template(
        'similarity_results.html',
        similarities=result_files,
        total_similarity=total_similarity,
        highlighted_pdf=output_pdf_filename
    )

@app.route('/check_similarity_from_upload')
@role_required(['admin'])
def check_similarity_from_upload():
    file_path = request.args.get('file_path')
    compare_to = request.args.get('compare_to', 'skripsi')

    if not file_path or not os.path.exists(file_path):
        flash("File tidak ditemukan", "danger")
        return redirect('/upload')

    return process_similarity_from_path(file_path, compare_to)

def get_file_path(history_id=None, user_id=None, doc_id=None, file_name=None):
    """
    Mengambil path file dari database berdasarkan history_id atau kombinasi user_id & doc_id.

    Args:
        history_id (int, optional): ID dari history di database.
        user_id (int, optional): ID pengguna.
        doc_id (int, optional): ID dokumen.
        file_name (str, optional): Nama file yang di-upload.

    Returns:
        str: Path ke file PDF atau None jika tidak ditemukan.
    """
    # Menggunakan try-except untuk menangani error koneksi dan query
    try:
        # Koneksi ke database
        conn = psycopg2.connect(
            "dbname=db_similarity user=postgres password=admin host=localhost"
        )
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

        # Query untuk mencari berdasarkan parameter yang valid
        if history_id:
            # Cari berdasarkan history_id langsung
            query = "SELECT result_file_path FROM history WHERE id = %s"
            cur.execute(query, (history_id,))
        elif user_id and doc_id and file_name:
            # Cari berdasarkan kombinasi user_id, doc_id, dan file_name
            query = """
                SELECT result_file_path FROM history 
                WHERE user_id = %s AND document_id = %s AND uploaded_file_name = %s
                ORDER BY id DESC LIMIT 1
            """
            cur.execute(query, (user_id, doc_id, file_name))
        else:
            # Jika tidak ada parameter yang valid
            return None

        # Mendapatkan hasil query
        result = cur.fetchone()

        # Menutup koneksi
        cur.close()
        conn.close()

        # Mengembalikan path file jika ditemukan, atau None jika tidak
        return result[0] if result else None

    except (psycopg2.Error, Exception) as e:
        # Menangani jika terjadi error
        print(f"Error occurred while fetching file path: {e}")
        return None


@app.route('/admin/view/<int:user_id>/<int:history_id>')
def admin_view_result(user_id, history_id):

    conn = psycopg2.connect(
        "dbname=db_similarity user=postgres password=admin host=localhost"
    )
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT result_file_path FROM history WHERE id = %s", (history_id,))
    result = cur.fetchone()
    cur.close()
    conn.close()
    print("Result from DB:", result)

    go_to_path = os.path.join(app.config['UPLOAD_FOLDER'], 'skripsi_highlights', str(user_id), result['result_file_path'])
    print("File path to serve:", go_to_path)

    if go_to_path:
        return send_file(go_to_path, mimetype='application/pdf', as_attachment=False)
    abort(404, description="File tidak ditemukan.")

@app.route('/user/highlight-exum/<int:user_id>/<filename>')
def serve_highlight_exum(user_id, filename):
    # Menggunakan os.path.join dan mengganti backslash dengan slash agar path benar
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'results', str(user_id), filename.replace('\\', '/')) 
    
    # Menampilkan error 404 jika file tidak ditemukan
    if os.path.exists(file_path):
        return send_file(file_path, mimetype='application/pdf', as_attachment=False)
    else:
        abort(404, description="File tidak ditemukan.")  # Menampilkan error jika file tidak ada

@app.route('/user/highlight-skripsi/<int:user_id>/<filename>')
def serve_highlight_skripsi(user_id, filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'skripsi_highlights', str(user_id), filename)
    if os.path.exists(file_path):
        return send_file(file_path, mimetype='application/pdf', as_attachment=False)
    else:
        return "File tidak ditemukan.", 404



@app.route('/some_route')
def some_route():
    # Data yang ingin di-zip
    list1 = ['a', 'b', 'c']
    list2 = [1, 2, 3]
    
    # Menambahkan zip ke dalam konteks template
    return render_template('some_template.html', list1=list1, list2=list2, zip=zip)

@app.route('/user/view/<int:history_id>')
def user_view_result(history_id):
    conn = psycopg2.connect(
        "dbname=db_similarity user=postgres password=admin host=localhost"
    )
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT result_file_path FROM history WHERE id = %s", (history_id,))
    result = cur.fetchone()
    cur.close()
    conn.close()

    if result and result[0] and os.path.exists(result[0]):
        return send_file(result[0], mimetype='application/pdf', as_attachment=False)
    abort(404, description="File tidak ditemukan.")


@app.route('/admin/view-original/<filename>')
@role_required(['admin'])
def view_admin_document(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'admin_documents', filename)

    if os.path.exists(file_path):
        return send_file(file_path, mimetype='application/pdf', as_attachment=False)
    
    abort(404, description="File asli tidak ditemukan.")


@app.route('/admin/delete-document/<int:doc_id>')
@role_required(['admin'])
def admin_delete_document(doc_id):
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    cur.execute("SELECT file_name FROM documents WHERE id = %s", (doc_id,))
    result = cur.fetchone()


    if result:
        filename = result['file_name']
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'admin_documents', filename)

        print("DEBUG: Akan hapus file", file_path)
        print("DEBUG: File exists:", os.path.exists(file_path))

        if os.path.exists(file_path):
            os.remove(file_path)
            print("DEBUG: File dihapus dari sistem.")
        else:
            print("DEBUG: File tidak ditemukan!")

    # Hapus dari database
    cur.execute("DELETE FROM documents WHERE id = %s", (doc_id,))
    print("DEBUG: Query DELETE dijalankan untuk dokumen ID", doc_id)

    conn.commit()
    cur.close()
    conn.close()

    flash("Dokumen berhasil dihapus", "success")
    return redirect('/admin/documents')




@app.route('/admin/edit-document/<int:doc_id>', methods=['GET', 'POST'])
@role_required(['admin'])
def edit_document(doc_id):
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    if request.method == 'POST':
        new_title = request.form['title']
        cur.execute("UPDATE documents SET title = %s WHERE id = %s", (new_title, doc_id))
        conn.commit()
        flash("Judul dokumen berhasil diperbarui.", "success")
        cur.close()
        conn.close()
        return redirect('/admin/documents')


    cur.execute("SELECT id, title FROM documents WHERE id = %s", (doc_id,))
    document = cur.fetchone()
    cur.close()
    conn.close()

    return render_template('admin/edit_document.html', document=document)


if __name__ == '__main__':
    init_db()
    print(app.url_map)

    # Tambahan: Tes similarity manual (optional)
    try:
        result = test_similarity(model, 1)
        print("Dokumen yang mirip dengan ID 1:", result)
    except Exception as e:
        print("Gagal melakukan pengecekan similarity:", e)

    # Jalankan aplikasi Flask
    app.run(debug=True)
